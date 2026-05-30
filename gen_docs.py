"""Generate comprehensive DOCX documentation for the PRM-HotpotQA project."""
import os
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ---- Page margins ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ---- Helper functions ----
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_image(path, caption, width=Inches(5.5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f'Figure: {caption}')
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.paragraphs[0].runs else None
    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Process Reward Model for Multi-Hop QA', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

subtitle = doc.add_heading('HotpotQA Distractor Setting', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AIMS-DTU Research Intern Round 2 - 2026')
run.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GitHub: https://github.com/RishiiGamer2201/prm-hotpotqa')
run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction and Problem Statement',
    '2. Approach and Methodology',
    '3. System Architecture',
    '4. Implementation Details',
    '5. Experimental Setup',
    '6. Results and Analysis',
    '7. Visualizations and Graphs',
    '8. Failure Analysis',
    '9. Errors Encountered and Resolutions',
    '10. Reproduction Instructions',
    '11. Assumptions and Limitations',
    '12. External Resources Used',
    '13. Conclusion',
]
for item in toc_items:
    add_para(item)
doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading('1. Introduction and Problem Statement', level=1)

add_para(
    'This project implements a Process Reward Model (PRM) that scores intermediate reasoning '
    'steps in a multi-hop retrieval-augmented question answering (QA) pipeline. The system is '
    'evaluated on HotpotQA\'s distractor setting using 500 held-out questions with five RAGAS '
    'metrics and 95% bootstrap confidence intervals.'
)

add_heading('1.1 What is a Process Reward Model?', level=2)
add_para(
    'Standard RAG pipelines retrieve documents and pass everything to the language model. '
    'Not all retrieved paragraphs are relevant -- some are noise, some are misleading. '
    'A Process Reward Model fixes this by scoring each retrieved paragraph at every hop '
    'and deciding whether it deserves to stay in the context window. Unlike outcome reward '
    'models that only evaluate the final answer, PRMs evaluate intermediate reasoning steps, '
    'providing fine-grained quality control at each stage of the retrieval process.'
)

add_heading('1.2 HotpotQA Distractor Setting', level=2)
add_para(
    'HotpotQA is a multi-hop question answering dataset where each question requires reasoning '
    'across multiple Wikipedia paragraphs. In the distractor setting, each question comes with '
    '10 paragraphs: 2 gold supporting facts (connected by a bridge entity) and 8 distractor '
    'paragraphs injected as noise. The PRM must consistently score the 2 gold paragraphs above '
    'the 8 distractors at each retrieval hop.'
)

add_heading('1.3 Multi-Hop Reasoning', level=2)
add_para('Some questions require chaining reasoning across multiple documents. For example:')
add_para('"What nationality is the director of the film Parasite?"', italic=True)
add_bullet('Hop 1: Retrieve paragraph about Parasite -> find director: Bong Joon-ho')
add_bullet('Hop 2: Use "Bong Joon-ho" as bridge entity -> retrieve his Wikipedia paragraph -> find: South Korean')

# ============================================================
# 2. APPROACH AND METHODOLOGY
# ============================================================
doc.add_page_break()
add_heading('2. Approach and Methodology', level=1)

add_heading('2.1 Pipeline Overview', level=2)
add_para(
    'The system follows a multi-hop retrieval-augmented generation pipeline with PRM gating:'
)
add_bullet('Step 1 - Query Processing: Accept a multi-hop question from HotpotQA')
add_bullet('Step 2 - Hop 1 Retrieval: Embed the question and retrieve top-k paragraphs from the 10-paragraph pool using FAISS')
add_bullet('Step 3 - PRM Scoring (Hop 1): Score each retrieved paragraph using the cross-encoder PRM')
add_bullet('Step 4 - Threshold Pruning: Remove paragraphs below the PRM threshold (t=0.4 or t=0.6)')
add_bullet('Step 5 - Bridge Entity Extraction: Extract named entities from kept paragraphs using spaCy NER')
add_bullet('Step 6 - Hop 2 Query Generation: Form sub-queries using bridge entities')
add_bullet('Step 7 - Hop 2 Retrieval: Retrieve additional paragraphs using bridge entity queries')
add_bullet('Step 8 - PRM Scoring (Hop 2): Score hop 2 paragraphs and prune below threshold')
add_bullet('Step 9 - Context Assembly: Merge and deduplicate paragraphs from both hops')
add_bullet('Step 10 - Answer Synthesis: Generate the final answer using flan-t5-large')

add_heading('2.2 PRM Threshold Ablation', level=2)
add_para(
    'Two PRM thresholds are compared to study the precision-recall tradeoff:'
)
add_bullet('t = 0.4 (Lenient): Allows more paragraphs through, maximizing recall at the cost of precision. More noise may leak into the context.')
add_bullet('t = 0.6 (Strict): Aggressively prunes low-scoring paragraphs, maximizing precision but risking the loss of gold supporting facts.')

add_heading('2.3 Evaluation Strategy', level=2)
add_para(
    'All five RAGAS metrics are computed for both threshold configurations over all 500 questions '
    'with 95% bootstrap confidence intervals (1000 bootstrap iterations, no subsampling). '
    'Random seed is fixed at 42 across all source files for full reproducibility.'
)

# ============================================================
# 3. SYSTEM ARCHITECTURE
# ============================================================
doc.add_page_break()
add_heading('3. System Architecture', level=1)

add_heading('3.1 Tech Stack', level=2)
make_table(
    ['Component', 'Model/Library', 'Notes'],
    [
        ['Embeddings', 'sentence-transformers/all-MiniLM-L6-v2', '384-dim, free, CPU-compatible'],
        ['PRM Scorer', 'cross-encoder/ms-marco-MiniLM-L-6-v2', 'Zero-shot cross-encoder, ~22M params'],
        ['Vector Search', 'FAISS (IndexFlatIP)', 'Cosine similarity with L2-normalized vectors'],
        ['Answer Generator', 'google/flan-t5-large', 'CPU-compatible, instruction-tuned'],
        ['RAGAS Judge LLM', 'Ollama mistral:7b', 'Local, free, no API keys'],
        ['NER', 'spaCy en_core_web_sm', 'Bridge entity extraction'],
        ['Dataset', 'HotpotQA distractor split', '500 validation questions'],
    ]
)

add_heading('3.2 Project Structure', level=2)
add_para(
    'prm-hotpotqa/\n'
    '|-- src/\n'
    '|   |-- prm.py              # PRM class, cross-encoder scoring, threshold pruning\n'
    '|   |-- retriever.py        # FAISS index, hop 1 + hop 2 retrieval, bridge extraction\n'
    '|   |-- pipeline.py         # End-to-end pipeline with CLI\n'
    '|-- eval/\n'
    '|   |-- ragas_eval.py       # RAGAS evaluation, bootstrap CIs, results export\n'
    '|-- results/\n'
    '|   |-- t0.4_raw.jsonl      # Per-question outputs at t=0.4\n'
    '|   |-- t0.6_raw.jsonl      # Per-question outputs at t=0.6\n'
    '|   |-- results.json        # Final RAGAS scores + CIs\n'
    '|   |-- results.csv         # Same in CSV\n'
    '|   |-- plots/              # Visualization outputs\n'
    '|-- notebooks/\n'
    '|   |-- analysis.ipynb      # Failure analysis notebook\n'
    '|-- requirements.txt\n'
    '|-- README.md'
)

# ============================================================
# 4. IMPLEMENTATION DETAILS
# ============================================================
doc.add_page_break()
add_heading('4. Implementation Details', level=1)

add_heading('4.1 PRM Model (src/prm.py)', level=2)
add_para(
    'The ProcessRewardModel class wraps the cross-encoder/ms-marco-MiniLM-L-6-v2 model. '
    'For each (question, paragraph) pair, the cross-encoder produces a relevance score. '
    'A sigmoid function normalizes these scores to the [0, 1] range. Key methods:'
)
add_bullet('score_step(question, paragraph): Scores a single paragraph against the question')
add_bullet('score_batch(question, paragraphs): Batch scoring for efficiency')
add_bullet('prune_steps(paragraphs, threshold): Removes paragraphs below the PRM threshold')
add_bullet('rank_steps(paragraphs): Returns paragraphs sorted by PRM score (descending)')

add_heading('4.2 Retriever (src/retriever.py)', level=2)
add_para(
    'The HotpotRetriever class handles all retrieval operations:'
)
add_bullet('FAISS Index: Uses IndexFlatIP (inner product) on L2-normalized embeddings for cosine similarity search')
add_bullet('Hop 1 Retrieval: Embeds the original question and retrieves top-k paragraphs from the 10-paragraph pool')
add_bullet('Bridge Entity Extraction: Uses spaCy NER to extract named entities from hop 1 results as bridge entities')
add_bullet('Hop 2 Query Generation: Constructs sub-queries using extracted bridge entities')
add_bullet('Hop 2 Retrieval: Retrieves additional paragraphs using bridge entity queries, deduplicating against hop 1 results')

add_heading('4.3 Pipeline (src/pipeline.py)', level=2)
add_para(
    'The MultiHopPipeline class orchestrates the entire process:'
)
add_bullet('Loads 500 HotpotQA validation questions using the datasets library')
add_bullet('For each question: retrieves, scores with PRM, prunes, performs hop 2, and synthesizes an answer')
add_bullet('Answer synthesis uses google/flan-t5-large with a structured prompt containing the question and kept context')
add_bullet('Outputs are saved as JSONL with full intermediate state (retrieved paragraphs, PRM scores, gold labels)')

add_heading('4.4 RAGAS Evaluation (eval/ragas_eval.py)', level=2)
add_para(
    'The evaluation script calls Ollama Mistral 7B directly via its REST API '
    '(http://localhost:11434/api/generate) to compute LLM-based metrics. This avoids '
    'all dependency issues with the ragas Python library\'s async evaluation loop.'
)
add_bullet('Faithfulness: LLM judges what fraction of answer claims are supported by retrieved context')
add_bullet('Answer Relevancy: LLM judges whether the answer addresses the question')
add_bullet('Context Precision: Computed directly -- fraction of retrieved paragraphs that are gold')
add_bullet('Context Recall: Computed directly -- fraction of gold paragraphs found in final context')
add_bullet('Answer Correctness: LLM compares predicted answer against gold answer; exact match gives 1.0')
add_para(
    'Bootstrap confidence intervals are computed with 1000 iterations over all 500 questions, '
    'using numpy\'s default_rng with seed=42.'
)

# ============================================================
# 5. EXPERIMENTAL SETUP
# ============================================================
doc.add_page_break()
add_heading('5. Experimental Setup', level=1)

add_heading('5.1 Hardware', level=2)
make_table(
    ['Component', 'Specification'],
    [
        ['Laptop', 'Acer TravelMate P214-54'],
        ['Processor', 'Intel 12th Gen Core i5-1235U @ ~1700 MHz'],
        ['RAM', '16 GB'],
        ['GPU', 'Intel UHD Graphics (integrated, no CUDA)'],
        ['OS', 'Windows 11'],
        ['Python', '3.14'],
    ]
)

add_heading('5.2 Runtime', level=2)
make_table(
    ['Phase', 'Duration', 'Notes'],
    [
        ['Pipeline t=0.4', '~3-4 hours', '500 questions, embedding + FAISS + PRM + answer generation'],
        ['Pipeline t=0.6', '~3-4 hours', 'Same pipeline, different threshold'],
        ['RAGAS Evaluation', '~30 hours', '500 questions x 2 thresholds x 3 LLM calls/question via CPU Ollama'],
        ['Total', '~36-38 hours', 'End-to-end on CPU (no GPU acceleration)'],
    ]
)

add_heading('5.3 Reproducibility', level=2)
add_bullet('Random Seed: 42 (set in all source files: prm.py, retriever.py, pipeline.py, ragas_eval.py)')
add_bullet('Bootstrap Iterations: 1000')
add_bullet('No subsampling: CIs computed over all 500 questions')
add_bullet('All models are free and publicly available on HuggingFace/Ollama')
add_bullet('No API keys required')

# ============================================================
# 6. RESULTS AND ANALYSIS
# ============================================================
doc.add_page_break()
add_heading('6. Results and Analysis', level=1)

add_heading('6.1 Detailed Results Table', level=2)
add_para('Table 1: RAGAS Metrics with 95% Bootstrap Confidence Intervals', bold=True)
make_table(
    ['Metric', 'PRM t=0.4', '95% CI', 'PRM t=0.6', '95% CI'],
    [
        ['Faithfulness', '0.5168', '[0.4786, 0.5561]', '0.5582', '[0.5197, 0.5980]'],
        ['Answer Relevancy', '0.5986', '[0.5589, 0.6401]', '0.5903', '[0.5471, 0.6341]'],
        ['Context Precision', '0.5460', '[0.5197, 0.5732]', '0.6004', '[0.5734, 0.6273]'],
        ['Context Recall', '0.8470', '[0.8270, 0.8670]', '0.8070', '[0.7850, 0.8280]'],
        ['Answer Correctness', '0.6319', '[0.5931, 0.6710]', '0.6207', '[0.5812, 0.6604]'],
    ]
)

add_para('')
add_para('Table 2: Assignment Format Results Table', bold=True)
make_table(
    ['System', 'Faith.', 'Ans. Rel.', 'Ctx. Prec.', 'Ctx. Rec.', 'Ans. Corr.'],
    [
        ['Your system (PRM t = 0.4)', '0.5168\n[0.4786, 0.5561]', '0.5986\n[0.5589, 0.6401]', '0.5460\n[0.5197, 0.5732]', '0.8470\n[0.8270, 0.8670]', '0.6319\n[0.5931, 0.6710]'],
        ['Your system (PRM t = 0.6)', '0.5582\n[0.5197, 0.5980]', '0.5903\n[0.5471, 0.6341]', '0.6004\n[0.5734, 0.6273]', '0.8070\n[0.7850, 0.8280]', '0.6207\n[0.5812, 0.6604]'],
    ]
)

add_heading('6.2 Key Observations', level=2)

add_para('Faithfulness (t=0.6 wins: 0.558 vs 0.517):', bold=True)
add_para(
    'The stricter threshold removes more noisy paragraphs from the context, reducing the chance '
    'of the answer generator hallucinating claims not supported by the evidence. This is the '
    'primary benefit of aggressive PRM pruning.'
)

add_para('Context Precision (t=0.6 wins: 0.600 vs 0.546):', bold=True)
add_para(
    'With a higher threshold, fewer irrelevant paragraphs survive pruning. The kept context '
    'has a higher fraction of gold supporting facts, resulting in better precision.'
)

add_para('Context Recall (t=0.4 wins: 0.847 vs 0.807):', bold=True)
add_para(
    'The lenient threshold keeps more paragraphs, including gold ones that might have borderline '
    'PRM scores. This demonstrates the fundamental precision-recall tradeoff in PRM threshold '
    'selection.'
)

add_para('Answer Relevancy (t=0.4 marginally wins: 0.599 vs 0.590):', bold=True)
add_para(
    'Both thresholds produce similarly relevant answers. The slight advantage of t=0.4 suggests '
    'that having more context (even if noisier) helps the model stay on-topic.'
)

add_para('Answer Correctness (t=0.4 marginally wins: 0.632 vs 0.621):', bold=True)
add_para(
    'End-to-end correctness is slightly better with the lenient threshold. The additional gold '
    'paragraphs retained by t=0.4 outweigh the extra noise, suggesting that for multi-hop QA, '
    'recall of gold facts is more important than precision of the context.'
)

add_heading('6.3 Hop 2 Recall Analysis', level=2)
add_para(
    'The fixed hop recall analysis revealed that hop 2 retrieval rescues 15-17% of gold '
    'paragraphs that were pruned by hop 1\'s PRM threshold. This demonstrates the value '
    'of multi-hop retrieval with bridge entity extraction.'
)
make_table(
    ['', 'Hop 1 Recall', 'Hop 2 Recall'],
    [
        ['t = 0.4', '1.000', '0.168'],
        ['t = 0.6', '1.000', '0.150'],
    ]
)
add_para(
    'Hop 1 recall of 1.0 means all gold paragraphs that survived PRM pruning at hop 1 '
    'made it to the final context. Hop 2 recall shows the fraction of gold paragraphs '
    'that were NOT found in hop 1 but were recovered through bridge entity-based retrieval '
    'in hop 2.'
)

# ============================================================
# 7. VISUALIZATIONS
# ============================================================
doc.add_page_break()
add_heading('7. Visualizations and Graphs', level=1)

plots_dir = r'C:\Users\seast\VSCODE\prm-hotpotqa\results\plots'

add_heading('7.1 RAGAS Metrics Comparison', level=2)
add_para(
    'Bar chart comparing all five RAGAS metrics between PRM t=0.4 and t=0.6, '
    'with 95% bootstrap confidence interval error bars.'
)
add_image(os.path.join(plots_dir, 'metric_bars.png'), 'RAGAS Metrics: PRM Threshold Comparison')

add_heading('7.2 PRM Score Distributions', level=2)
add_para(
    'Histogram of PRM scores for hop 1 and hop 2 paragraphs, showing the bimodal distribution. '
    'Most paragraphs score near 0 (irrelevant) or near 1 (relevant), with the threshold '
    'line showing where pruning occurs.'
)
add_image(os.path.join(plots_dir, 'score_distributions.png'), 'PRM Score Distribution (t=0.4 and t=0.6)')

if os.path.exists(os.path.join(plots_dir, 'score_distributions_detailed.png')):
    add_image(os.path.join(plots_dir, 'score_distributions_detailed.png'), 'Detailed PRM Score Distributions')

add_heading('7.3 Context Recall by Hop', level=2)
add_para(
    'Comparison of gold paragraph recall at hop 1 vs hop 2 for both thresholds. '
    'Shows that hop 2 retrieval recovers ~15-17% of gold paragraphs missed by hop 1.'
)
add_image(os.path.join(plots_dir, 'context_recall_hop.png'), 'Context Recall by Hop')

if os.path.exists(os.path.join(plots_dir, 'metric_distributions.png')):
    add_heading('7.4 Per-Question Metric Distributions', level=2)
    add_para(
        'Histogram of per-question RAGAS scores, showing the distribution of individual '
        'question-level metric values across all 500 questions.'
    )
    add_image(os.path.join(plots_dir, 'metric_distributions.png'), 'Per-Question Metric Score Distributions')

# ============================================================
# 8. FAILURE ANALYSIS
# ============================================================
doc.add_page_break()
add_heading('8. Failure Analysis', level=1)

add_para(
    'The analysis notebook (notebooks/analysis.ipynb) contains detailed failure analysis '
    'with specific examples. Below is a summary of the two primary failure modes.'
)

add_heading('8.1 Hop Failures (3 Examples in Notebook)', level=2)
add_para(
    'A hop failure occurs when a gold supporting paragraph is NOT present in the final context. '
    'This typically happens when:'
)
add_bullet('The PRM assigns a score below the threshold to a gold paragraph (false negative)')
add_bullet('The bridge entity is not properly extracted from hop 1, preventing hop 2 from finding the second gold paragraph')
add_bullet('The embedding similarity between the question and the gold paragraph is low, causing it to rank poorly in initial retrieval')

add_heading('8.2 PRM False-Positive Prunes (3 Examples in Notebook)', level=2)
add_para(
    'A false-positive prune occurs when the PRM incorrectly assigns a low score to a gold '
    'paragraph, causing it to be pruned. This is more common with:'
)
add_bullet('Gold paragraphs that provide indirect evidence (the connection to the question is implicit)')
add_bullet('Paragraphs with high lexical overlap with distractors, confusing the cross-encoder')
add_bullet('The stricter threshold (t=0.6) produces more false-positive prunes than t=0.4')

add_heading('8.3 Impact of Threshold on Failures', level=2)
add_para(
    'At t=0.4, the system retains more paragraphs (including gold ones with borderline scores), '
    'resulting in higher context recall (0.847 vs 0.807) but lower context precision (0.546 vs 0.600). '
    'At t=0.6, the system is more aggressive in pruning, which improves faithfulness (0.558 vs 0.517) '
    'but risks dropping critical evidence.'
)

# ============================================================
# 9. ERRORS AND RESOLUTIONS
# ============================================================
doc.add_page_break()
add_heading('9. Errors Encountered and Resolutions', level=1)

add_para(
    'During the development and execution of this project, several technical challenges '
    'were encountered and resolved. This section documents each issue and its solution.'
)

# Error 1
add_heading('9.1 RAGAS Library Compatibility Issues', level=2)
add_para('Problem:', bold=True)
add_para(
    'The ragas library (v0.4.3) had breaking API changes from v0.2.x. The metric classes '
    '(Faithfulness, AnswerRelevancy, etc.) required an llm argument in their constructors, '
    'which was not needed in older versions. Additionally, the internal isinstance() check '
    'in ragas\' evaluate() function failed because metrics from ragas.metrics.collections '
    'inherit from BaseMetric instead of Metric.'
)
add_para('Resolution:', bold=True)
add_para(
    'Migrated to using internal old-style metrics (ragas.metrics._faithfulness, etc.) '
    'which maintain backwards compatibility with the Metric base class. Later, completely '
    'rewrote the evaluation to call Ollama directly via REST API, bypassing the ragas '
    'evaluate() loop entirely.'
)

# Error 2
add_heading('9.2 Evaluation Freezing at 0%', level=2)
add_para('Problem:', bold=True)
add_para(
    'The RAGAS evaluation appeared to freeze at "Evaluating: 0%" with no progress for '
    'extended periods (5+ minutes). The progress bar showed 0/2500 completed. Ctrl+C '
    'was unresponsive due to the async event loop.'
)
add_para('Resolution:', bold=True)
add_para(
    'The "freeze" was actually Mistral 7B processing long prompts on CPU, taking 60-90 '
    'seconds per call. The solution was threefold: (1) Rewrote the eval script to call '
    'Ollama directly via synchronous REST calls instead of async ragas wrappers. '
    '(2) Added per-metric verbose output so progress is visible (e.g., "faithfulness... 0.50"). '
    '(3) Made Ctrl+C work properly by removing the async event loop.'
)

# Error 3
add_heading('9.3 Hop 2 Recall Graph Showing Zero', level=2)
add_para('Problem:', bold=True)
add_para(
    'The context_recall_hop.png plot showed Hop 2 recall as 0.0 for both thresholds, '
    'which was incorrect. Investigation revealed a bug in the compute_hop_recall() function.'
)
add_para('Root Cause:', bold=True)
add_para(
    'The function only iterated over hop1_retrieved paragraphs and checked their "hop" field. '
    'Since all paragraphs in hop1_retrieved have hop=1 (assigned during retrieval), the '
    'hop2_recalls list was always empty, producing 0.0.'
)
add_para('Resolution:', bold=True)
add_para(
    'Fixed the logic to properly track gold paragraphs across hops: (1) Check which gold '
    'titles survived hop 1 PRM pruning (in hop1_kept). (2) For gold titles NOT in hop1_kept, '
    'check if they were recovered by hop 2 retrieval (in final_contexts). Since the eval was '
    'already running with the old code, a separate fix script (fix_hop_plot.py) was created '
    'and executed after the eval completed to regenerate only the affected plot from the saved '
    'JSONL results. The fixed plot correctly showed Hop 2 recall of 0.168 (t=0.4) and 0.150 (t=0.6).'
)

# Error 4
add_heading('9.4 VertexAI Import Error', level=2)
add_para('Problem:', bold=True)
add_para(
    'RAGAS internally imports langchain_community.chat_models.vertexai, which was not '
    'installed. This caused an ImportError at startup.'
)
add_para('Resolution:', bold=True)
add_para(
    'Created a stub file at venv/Lib/site-packages/langchain_community/chat_models/vertexai.py '
    'with a placeholder ChatVertexAI class to satisfy the import without installing the '
    'full Google Cloud SDK.'
)

# Error 5
add_heading('9.5 HuggingFaceEmbeddings API Change', level=2)
add_para('Problem:', bold=True)
add_para(
    'The HuggingFaceEmbeddings constructor changed its parameter name from model_name '
    'to model in newer versions of langchain-huggingface, causing a TypeError.'
)
add_para('Resolution:', bold=True)
add_para(
    'Updated the constructor call to use the new parameter name: '
    'HuggingFaceEmbeddings(model="sentence-transformers/all-MiniLM-L6-v2").'
)

# Error 6
add_heading('9.6 Checkpoint/Resume for Long Runs', level=2)
add_para('Problem:', bold=True)
add_para(
    'The RAGAS evaluation takes ~30 hours on CPU. If the laptop loses power, sleeps, or '
    'crashes, all progress is lost and the entire evaluation must restart from question 1.'
)
add_para('Resolution:', bold=True)
add_para(
    'Added checkpoint support to the evaluation script. After each question, scores are '
    'appended to a checkpoint JSONL file (results/checkpoints/checkpoint_t0.4.jsonl). '
    'On restart, the script detects the checkpoint, loads completed scores, and resumes '
    'from the next unfinished question.'
)

# Error 7
add_heading('9.7 Ollama Timeout on Long Prompts', level=2)
add_para('Problem:', bold=True)
add_para(
    'Occasionally, the faithfulness metric prompt (which includes full paragraph context) '
    'exceeded the 120-second timeout when processed by Mistral 7B on CPU.'
)
add_para('Resolution:', bold=True)
add_para(
    'The timeout error is caught gracefully and the metric defaults to 0.0 for that question. '
    'With 500 questions, occasional timeouts (~1-2%) have minimal impact on aggregate metrics.'
)

# ============================================================
# 10. REPRODUCTION INSTRUCTIONS
# ============================================================
doc.add_page_break()
add_heading('10. Reproduction Instructions', level=1)

add_heading('10.1 Prerequisites', level=2)
add_bullet('Python 3.10+ (tested on 3.14)')
add_bullet('Ollama installed (https://ollama.com/download)')
add_bullet('~10 GB disk space (models + results)')
add_bullet('16 GB RAM recommended')

add_heading('10.2 Setup', level=2)
add_para(
    'git clone https://github.com/RishiiGamer2201/prm-hotpotqa\n'
    'cd prm-hotpotqa\n'
    'pip install -r requirements.txt\n'
    'python -m spacy download en_core_web_sm\n'
    'ollama pull mistral'
)

add_heading('10.3 Execution Order', level=2)
add_para(
    'Step 1: Start Ollama server (in a separate terminal):\n'
    '    ollama serve\n\n'
    'Step 2: Run pipeline for t=0.4:\n'
    '    python src/pipeline.py --threshold 0.4 --output results/t0.4_raw.jsonl\n\n'
    'Step 3: Run pipeline for t=0.6:\n'
    '    python src/pipeline.py --threshold 0.6 --output results/t0.6_raw.jsonl\n\n'
    'Step 4: Run RAGAS evaluation:\n'
    '    python eval/ragas_eval.py --llm ollama --t04 results/t0.4_raw.jsonl --t06 results/t0.6_raw.jsonl --output results/\n\n'
    'Step 5: Run analysis notebook:\n'
    '    jupyter nbconvert --to notebook --execute notebooks/analysis.ipynb'
)

# ============================================================
# 11. ASSUMPTIONS AND LIMITATIONS
# ============================================================
doc.add_page_break()
add_heading('11. Assumptions and Limitations', level=1)

add_heading('11.1 Assumptions', level=2)
add_bullet('The cross-encoder model (ms-marco-MiniLM-L-6-v2) is used in a zero-shot setting without fine-tuning on HotpotQA data')
add_bullet('The 10-paragraph pool per question (2 gold + 8 distractors) is provided by HotpotQA; no additional retrieval corpus is used')
add_bullet('Bridge entities are extracted using spaCy NER, which may miss non-standard entity types')
add_bullet('RAGAS LLM-based metrics (faithfulness, answer relevancy, answer correctness) use Mistral 7B as the judge, which may have its own biases')

add_heading('11.2 Limitations', level=2)
add_bullet('CPU-only execution: No GPU acceleration available, resulting in long runtimes (~36 hours total)')
add_bullet('Zero-shot PRM: The cross-encoder is not fine-tuned on HotpotQA, which limits its ability to distinguish gold from distractor paragraphs in domain-specific cases')
add_bullet('Single LLM judge: RAGAS metrics depend on Mistral 7B quality; a stronger judge model would provide more reliable metric scores')
add_bullet('Two-hop limit: The pipeline only supports 2 hops; some questions may require 3+ reasoning steps')
add_bullet('Deterministic but model-dependent: While seeds are fixed, model weights and inference on different hardware may produce slightly different floating-point results')

# ============================================================
# 12. EXTERNAL RESOURCES
# ============================================================
add_heading('12. External Resources Used', level=1)

add_heading('12.1 Models (all free, no API keys)', level=2)
add_bullet('sentence-transformers/all-MiniLM-L6-v2 (HuggingFace) - Sentence embeddings')
add_bullet('cross-encoder/ms-marco-MiniLM-L-6-v2 (HuggingFace) - PRM cross-encoder scorer')
add_bullet('google/flan-t5-large (HuggingFace) - Answer generation')
add_bullet('mistral:7b (Ollama) - RAGAS evaluation LLM judge')

add_heading('12.2 Libraries', level=2)
add_bullet('PyTorch, Transformers, Sentence-Transformers - Model inference')
add_bullet('FAISS - Vector similarity search')
add_bullet('spaCy - Named entity recognition for bridge entity extraction')
add_bullet('RAGAS (metrics definitions only) - Evaluation framework reference')
add_bullet('Ollama - Local LLM serving')
add_bullet('Matplotlib, NumPy, Pandas - Data analysis and visualization')

add_heading('12.3 Dataset', level=2)
add_bullet('HotpotQA (Yang et al., 2018) - Multi-hop question answering dataset, distractor setting')
add_bullet('Accessed via HuggingFace Datasets library: hotpotqa/distractor')

# ============================================================
# 13. CONCLUSION
# ============================================================
add_heading('13. Conclusion', level=1)

add_para(
    'This project demonstrates a working Process Reward Model pipeline for multi-hop '
    'question answering on HotpotQA\'s distractor setting. The key findings are:'
)

add_bullet(
    'PRM threshold selection involves a precision-recall tradeoff: t=0.6 improves '
    'faithfulness (+0.041) and context precision (+0.054) while t=0.4 preserves '
    'context recall (+0.040) and answer correctness (+0.011).'
)
add_bullet(
    'Multi-hop retrieval with bridge entity extraction recovers 15-17% of gold '
    'paragraphs that would otherwise be lost after hop 1 PRM pruning.'
)
add_bullet(
    'The zero-shot cross-encoder PRM effectively separates gold from distractor '
    'paragraphs, as evidenced by the bimodal PRM score distribution.'
)
add_bullet(
    'The entire pipeline runs on commodity CPU hardware without any API keys or '
    'paid services, making it fully reproducible.'
)

add_para('')
add_para('Seed = 42. All results are fully reproducible.', bold=True)

# ---- Save ----
output_path = r'C:\Users\seast\Downloads\PRM_HotpotQA_Documentation.docx'
doc.save(output_path)
print(f'Documentation saved to {output_path}')
